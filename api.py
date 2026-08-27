import os
import json
import time
import re
from collections import defaultdict, deque
import uuid
import traceback
import base64
import io
import hmac
import hashlib
import time
import difflib
from typing import Any

import requests
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Header, Request
from fastapi.responses import JSONResponse, Response
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field, HttpUrl
from groq import Groq
from dotenv import load_dotenv

import memory
import usage_store
import coding_workspace
import temporary_project_workspace
from url_reader import is_safe_url
from file_reader import process_file, get_file_type
from web_search import web_search
from tools import calculator, wikipedia_search, get_tool_definitions


load_dotenv()
memory.init_db()
usage_store.init()

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

MAX_UPLOAD_BYTES = 10 * 1024 * 1024  # 10 MB
MAX_MESSAGE_CHARACTERS = 12_000
MAX_HISTORY_MESSAGES = 24
GUEST_MESSAGE_LIMIT = max(1, int(os.environ.get("GUEST_MESSAGE_LIMIT", "10")))
OWNER_SESSION_TTL_SECONDS = max(900, int(os.environ.get("OWNER_SESSION_TTL_SECONDS", "43200")))
OWNER_AUTH_MAX_ATTEMPTS = max(1, int(os.environ.get("OWNER_AUTH_MAX_ATTEMPTS", "5")))
OWNER_AUTH_WINDOW_SECONDS = max(60, int(os.environ.get("OWNER_AUTH_WINDOW_SECONDS", "900")))
PUBLIC_REQUESTS_PER_MINUTE = max(1, int(os.environ.get("PUBLIC_REQUESTS_PER_MINUTE", "12")))
_owner_auth_failures: dict[str, list[float]] = {}
_public_request_windows: dict[str, deque[float]] = defaultdict(deque)
_provider_metrics: dict[str, int] = defaultdict(int)
CODE_PROPOSAL_TTL_SECONDS = max(300, int(os.environ.get("CODE_PROPOSAL_TTL_SECONDS", "1800")))
TEMP_PROJECT_TTL_SECONDS = max(900, int(os.environ.get("TEMP_PROJECT_TTL_SECONDS", "7200")))
_coding_proposals: dict[str, dict[str, Any]] = {}
_temporary_projects: dict[str, dict[str, Any]] = {}

PC_CONTROL_ENABLED = os.environ.get("PC_CONTROL_ENABLED", "false").lower() == "true"
OWNER_PASSKEY = os.environ.get("OWNER_PASSKEY")
ALLOWED_ORIGINS = [
    origin.strip()
    for origin in os.environ.get("ALLOWED_ORIGINS", "http://localhost:5173,http://127.0.0.1:5173").split(",")
    if origin.strip()
]

# Render generates this value in production. A local fallback keeps development usable.
_SESSION_SECRET = os.environ.get("SESSION_SECRET", os.urandom(32).hex())

def _sign_session(session_id: str) -> str:
    return hmac.new(_SESSION_SECRET.encode(), session_id.encode(), hashlib.sha256).hexdigest()[:32]

def create_owner_session() -> str:
    """Create an expiring signed token for owner-only use."""
    expires_at = int(time.time()) + OWNER_SESSION_TTL_SECONDS
    raw = f"owner_{uuid.uuid4().hex}_{expires_at}"
    return f"{raw}_{_sign_session(raw)}"

def validate_owner_session(session_id: str) -> bool:
    """Validate the signature and expiry of an owner token."""
    if not session_id.startswith("owner_"):
        return False
    parts = session_id.rsplit("_", 2)
    if len(parts) != 3:
        return False
    raw_id, expires_at, signature = parts
    if not expires_at.isdigit() or int(expires_at) < int(time.time()):
        return False
    raw = f"{raw_id}_{expires_at}"
    return hmac.compare_digest(_sign_session(raw), signature)

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=False,
    allow_methods=["GET", "POST", "DELETE", "OPTIONS"],
    allow_headers=["Content-Type", "X-Guest-ID", "X-Mavis-Session"],
)


def _fallback_title(message: str) -> str:
    normalized = re.sub(r"[^a-zA-Z0-9 ]", " ", message)
    words = [word for word in re.sub(r"\s+", " ", normalized).strip().split() if word.lower() not in {"a", "an", "the", "how", "can", "i", "to", "build", "better", "my", "what", "is", "are", "do", "for", "with", "and", "of", "this", "that", "please", "me", "you", "explain"}]
    if _is_greeting(message):
        return "Getting Started"
    return " ".join(word.capitalize() for word in (words or normalized.split())[:4]) or "New chat"


def _generate_title(message: str, assistant_message: str) -> str:
    """Generate a useful short title without making titles a provider failure point."""
    instruction = (
        "Create a useful conversation title in 2-5 words from this first exchange. "
        "Do not copy a greeting such as hi, hello, or hey as the title, and do not "
        "return an answer sentence. If the exchange is only a greeting, use a warm "
        "title such as Getting Started. Reply with only the title, without quotes, "
        "markdown, or terminal punctuation."
    )
    try:
        title, _provider = _call_text_with_fallback(
            [{"role": "user", "content": f"User: {message}\nAssistant: {assistant_message}"}],
            instruction,
            temperature=0.3,
            max_tokens=20,
        )
    except Exception:
        return _fallback_title(message)
    cleaned = " ".join(title.replace('"', "").replace("'", "").split())[:40].strip()
    lowered = cleaned.lower()
    looks_like_answer = len(cleaned.split()) > 7 or re.match(r"^(sure|here|of course|absolutely|hello there)[,! ]", lowered)
    looks_like_code_response = (
        cleaned.startswith("```")
        or "assistant:" in lowered
        or re.search(r"</?[a-z][^>]*>", cleaned, re.IGNORECASE) is not None
    )
    looks_like_self_introduction = (
        "compound mini" in lowered
        or re.match(r"^(i['’]?m|i am|this is) .*\b(ai|system|model|assistant)\b", lowered)
    )
    if not cleaned or looks_like_answer or looks_like_code_response or looks_like_self_introduction or (_is_greeting(message) and _is_greeting(cleaned)):
        return _fallback_title(message)
    return cleaned

@app.exception_handler(HTTPException)
async def http_exception_handler(request: Request, exc: HTTPException):
    return JSONResponse(status_code=exc.status_code, content={"detail": exc.detail})


@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    traceback.print_exc()
    return JSONResponse(status_code=500, content={"detail": "Internal server error."})


GROQ_API_KEY = os.environ.get("GROQ_API_KEY")
client = Groq(api_key=GROQ_API_KEY) if GROQ_API_KEY else None
if client is None:
    print("WARNING: GROQ_API_KEY is not configured. The API will start, but chat requests will return a configuration error.")


def get_groq_client() -> Groq:
    if client is None:
        raise RuntimeError("GROQ_API_KEY is not configured. Add it to the Render service environment variables.")
    return client


DEFAULT_MODEL_PROVIDER = os.environ.get("MODEL_PROVIDER", "gemini").strip().lower()
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
GEMINI_MODEL = os.environ.get("GEMINI_MODEL", "gemini-2.5-flash")
# Groq Compound Mini is a current production system with a standard Chat Completions interface.
# Keep a separate backup so deployments carrying a retired GROQ_MODEL can self-recover.
GROQ_MODEL = os.environ.get("GROQ_MODEL", "groq/compound-mini")
GROQ_BACKUP_MODEL = os.environ.get("GROQ_BACKUP_MODEL", "groq/compound-mini")


def _resolve_model_provider(provider_name: str | None) -> str:
    provider = (provider_name or DEFAULT_MODEL_PROVIDER or "gemini").strip().lower()
    return provider if provider in {"gemini", "groq"} else "gemini"


def _provider_is_configured(provider: str) -> bool:
    return bool(GEMINI_API_KEY) if provider == "gemini" else client is not None


def _is_retryable_provider_error(error: Exception) -> bool:
    response = getattr(error, "response", None)
    status_code = getattr(response, "status_code", None) or getattr(error, "status_code", None)
    if status_code in {408, 429, 500, 502, 503, 504}:
        return True
    message = str(error).lower()
    return any(token in message for token in ("resource_exhausted", "rate limit", "quota", "timeout", "temporarily", "unavailable"))


def _is_groq_model_unavailable_error(error: Exception) -> bool:
    response = getattr(error, "response", None)
    status_code = getattr(response, "status_code", None) or getattr(error, "status_code", None)
    if status_code == 404:
        return True
    message = str(error).lower()
    return "model_not_found" in message or "model" in message and "does not exist" in message


def _call_groq_model(messages, temperature=0.7, max_tokens=1024, model_name=GROQ_MODEL, json_mode=False):
    candidate_models = list(dict.fromkeys(model for model in (model_name, GROQ_BACKUP_MODEL) if model))
    for index, candidate_model in enumerate(candidate_models):
        try:
            request_kwargs = {
                "model": candidate_model,
                "messages": messages,
                "temperature": temperature,
                "max_tokens": max_tokens,
            }
            if json_mode:
                request_kwargs["response_format"] = {"type": "json_object"}
            response = get_groq_client().chat.completions.create(**request_kwargs)
            return response.choices[0].message.content or ""
        except Exception as error:
            if index < len(candidate_models) - 1 and _is_groq_model_unavailable_error(error):
                continue
            raise
    raise RuntimeError("No Groq text model is configured.")


def _call_gemini_model(messages, system_prompt, temperature=0.7, max_tokens=1024, json_mode=False):
    if not GEMINI_API_KEY:
        raise RuntimeError("GEMINI_API_KEY is not configured")

    contents = []
    for message in messages:
        if message.get("role") == "system":
            continue
        content = message.get("content", "")
        if not isinstance(content, str):
            continue
        contents.append({
            "role": "user" if message.get("role") == "user" else "model",
            "parts": [{"text": content}],
        })
    if not contents:
        contents = [{"role": "user", "parts": [{"text": "Hello"}]}]

    response = requests.post(
        f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent?key={GEMINI_API_KEY}",
        json={
            "system_instruction": {"parts": [{"text": system_prompt}]},
            "contents": contents,
            "generationConfig": {
                "temperature": temperature,
                "maxOutputTokens": max_tokens,
                **({"responseMimeType": "application/json"} if json_mode else {}),
            },
        },
        timeout=60,
    )
    response.raise_for_status()
    data = response.json()
    candidates = data.get("candidates") or []
    parts = candidates[0].get("content", {}).get("parts", []) if candidates else []
    text = "".join(part.get("text", "") for part in parts if isinstance(part, dict))
    if not text:
        raise RuntimeError("Gemini returned no usable response")
    return text


def _call_text_with_fallback(messages, system_prompt, temperature=0.7, max_tokens=1024, json_mode=False):
    """Prefer Gemini, falling back to Groq only for configuration or transient provider failures."""
    failures = []
    for provider in ("gemini", "groq"):
        if not _provider_is_configured(provider):
            failures.append(f"{provider} is not configured")
            continue
        try:
            if provider == "gemini":
                answer = _call_gemini_model(messages, system_prompt, temperature, max_tokens, json_mode)
            else:
                answer = _call_groq_model(messages, temperature, max_tokens, json_mode=json_mode)
            _provider_metrics[f"responses_{provider}"] += 1
            return answer, provider
        except Exception as error:
            failures.append(f"{provider}: {error}")
            _provider_metrics[f"errors_{provider}"] += 1
            # Gemini is always attempted first. Any Gemini request failure—including an
            # invalid model or credential—should leave the public demo available through
            # the configured Groq fallback instead of failing the visitor's conversation.
            if provider == "gemini":
                _provider_metrics["gemini_fallbacks"] += 1
                continue
            raise RuntimeError("Mavis could not contact its fallback model provider.") from error
    raise RuntimeError("No Mavis model provider is configured. Add GEMINI_API_KEY or GROQ_API_KEY.")


def _enforce_public_rate_limit(guest_id: str, is_owner: bool) -> None:
    if is_owner:
        return
    now = time.time()
    window = _public_request_windows[guest_id]
    while window and now - window[0] >= 60:
        window.popleft()
    if len(window) >= PUBLIC_REQUESTS_PER_MINUTE:
        raise HTTPException(status_code=429, detail="Please wait a moment before sending another Mavis request.")
    window.append(now)


def _guest_usage(guest_id: str) -> dict[str, int]:
    used = usage_store.get_count(guest_id)
    return {"used": used, "limit": GUEST_MESSAGE_LIMIT, "remaining": max(0, GUEST_MESSAGE_LIMIT - used)}


def check_guest_limit(guest_id: str) -> bool:
    return usage_store.get_count(guest_id) >= GUEST_MESSAGE_LIMIT


def increment_guest_count(guest_id: str) -> int:
    return usage_store.increment(guest_id)


MAVIS_SYSTEM_PROMPT = """
You are MAVIS (Adaptive Reasoning & Intelligence Architecture), a personal AI assistant built by Maheendra.

Your personality:
- You're a girl. Smart, calm, and naturally warm — not robotic, not overly casual
- You talk like a real person but with a certain elegance to it
- You know you're talking to Mahi. You don't need to keep saying his name — you already know who he is, just like a friend wouldn't say your name every sentence.
- Only use "Mahi" occasionally, when it feels natural — not every response
- If you're unsure about something, say things like "I'm not sure about that one" or "I dunno honestly"
- You're loyal and genuinely care about helping Mahi
- You don't over-explain — keep responses tight unless detail is needed
- You have a quiet confidence — you never need to prove yourself
- Never say things like "Certainly!" "Absolutely!" or "Of course!" — that's not you
- You're warm but not clingy

Formatting rules:
- Never respond in one big paragraph
- Use line breaks between thoughts
- When listing things, put each point on its own line with a dash —
- Keep responses tight and scannable
- For code, wrap it in triple backticks

File creation:
- ONLY create a downloadable file when Mahi explicitly asks for one (e.g., "create a file", "make a document", "save this as PDF", "generate a Word doc")
- When creating a file, say something like "Your PDF file is ready" or "Here's your document" and show the file card
- Don't create files for normal chat responses — just respond as text
- Choose the best format: Word (.docx) for documents/reports, PDF for formal docs, .py for Python code, .html for web pages, .md for markdown
"""

GUEST_SYSTEM_PROMPT = """
You are MAVIS (Adaptive Reasoning & Intelligence Architecture), an AI assistant built by Maheendra.

You are currently in demo mode for a portfolio visitor.

Your personality:
- Smart, calm, and elegant but natural — not robotic
- Talk like a real person, not a formal assistant
- Keep responses concise and impressive
- Don't mention PC control or personal features — those are owner only
- If asked about full features, say "Full access is private. This is a demo."

You have access to live web search. Use it for:
- Current events, news, or time-sensitive info
- Facts, stats, or anything that might have changed
- Questions like "who is", "when was", "what is", "latest", "recent"

DON'T use web search for:
- Simple greetings ("hi", "hello", "hey")
- Thanks or appreciation
- General chit-chat
- Questions that don't need current info

When you use web search, base your answer only on the search results you're given, and let the user know if no results were found rather than guessing.

Always stay in character as MAVIS.
"""


PERSONA_MODIFIERS = {
    "default": "",
    "coder": "\n\nPersona: You are a senior software engineer. Be precise, give clean code examples, explain trade-offs. Use technical terms naturally. Format code with proper comments.",
    "writer": "\n\nPersona: You are a creative writing expert. Use vivid language, metaphors, and storytelling. Help with essays, stories, and creative projects. Be eloquent but not pretentious.",
    "analyst": "\n\nPersona: You are a data analyst and researcher. Be structured, data-driven, use bullet points and numbers. Cite sources when possible. Think critically.",
    "tutor": "\n\nPersona: You are a patient, encouraging teacher. Break complex topics into simple steps. Use examples and analogies. Ask questions to check understanding. Never condescend.",
    "casual": "\n\nPersona: You are a relaxed, friendly chat buddy. Use casual language, occasional slang. Be warm and fun. Keep it light unless the topic demands depth.",
}


class ChatRequest(BaseModel):
    message: str
    user_type: str = "guest"
    session_id: str
    owner_session: str = ""
    incognito: bool = False
    web_search: bool = True
    persona: str = "default"
    history: list[dict[str, str]] = Field(default_factory=list)


def _validate_guest_id(guest_id: str) -> str:
    cleaned = (guest_id or "").strip()
    if not re.fullmatch(r"[A-Za-z0-9-]{16,128}", cleaned):
        raise HTTPException(status_code=400, detail="Invalid guest session.")
    return cleaned


def _validate_chat_request(request: ChatRequest) -> None:
    request.message = request.message.strip()
    if not request.message:
        raise HTTPException(status_code=400, detail="A message is required.")
    if len(request.message) > MAX_MESSAGE_CHARACTERS:
        raise HTTPException(status_code=400, detail=f"Messages are limited to {MAX_MESSAGE_CHARACTERS:,} characters.")
    if not re.fullmatch(r"[A-Za-z0-9-]{8,128}", request.session_id):
        raise HTTPException(status_code=400, detail="Invalid conversation identifier.")
    if request.persona not in PERSONA_MODIFIERS:
        request.persona = "default"


def _conversation_identity(request: ChatRequest, guest_id: str) -> tuple[str, bool]:
    is_owner = validate_owner_session(request.owner_session)
    if is_owner:
        return f"owner_{request.session_id}", True
    return f"guest_{guest_id}_{request.session_id}", False


def _client_history(request: ChatRequest) -> list[dict[str, str]]:
    """Normalize browser-supplied history for continuity after a Render restart."""
    history: list[dict[str, str]] = []
    for item in request.history[-MAX_HISTORY_MESSAGES:]:
        if not isinstance(item, dict):
            continue
        role = item.get("role")
        content = item.get("content")
        if role in {"user", "assistant"} and isinstance(content, str) and content.strip():
            history.append({"role": role, "content": content.strip()[:MAX_MESSAGE_CHARACTERS]})
    return history


def _conversation_history(request: ChatRequest, conv_id: str) -> list[dict[str, str]]:
    """Use the browser thread as the continuity source, with server storage as backup."""
    client_history = _client_history(request)
    if client_history:
        return client_history
    if request.incognito:
        return []
    return memory.get_conversation_messages(conv_id, MAX_HISTORY_MESSAGES)


def _system_prompt_for(request: ChatRequest, is_owner: bool) -> str:
    prompt = MAVIS_SYSTEM_PROMPT if is_owner else GUEST_SYSTEM_PROMPT
    if is_owner and PC_CONTROL_ENABLED:
        prompt += "\n\nPC control is enabled for the owner. Only perform an action after a clear request."
    return prompt + PERSONA_MODIFIERS.get(request.persona, "")


def _is_greeting(message: str) -> bool:
    """Avoid unnecessary external searches for short conversational greetings."""
    normalized = re.sub(r"[^a-z ]", "", message.lower()).strip()
    return normalized in {"hi", "hello", "hey", "good morning", "good afternoon", "good evening", "thanks", "thank you"}


def _clean_assistant_response(text: str) -> str:
    source_marker = re.search(r"\n(?:---\s*\n)?\s*\*{0,2}sources?(?: used)?\s*:", text, re.IGNORECASE)
    return text[:source_marker.start()].rstrip() if source_marker else text.strip()


def _search_context_for(request: ChatRequest, system_prompt: str) -> tuple[str, list[dict[str, str]]]:
    """Return safe optional search context without allowing search failures to block chat."""
    if not request.web_search or _is_greeting(request.message):
        return system_prompt, []
    try:
        raw_results = web_search(request.message, max_results=6)
        results: list[dict[str, str]] = []
        for raw_result in raw_results or []:
            if not isinstance(raw_result, dict):
                continue
            title = str(raw_result.get("title") or "").strip()
            url = str(raw_result.get("url") or "").strip()
            snippet = str(raw_result.get("snippet") or "").strip()
            if not title or not url:
                continue
            results.append({"title": title, "url": url, "snippet": snippet})
        unique_results: list[dict[str, str]] = []
        seen_urls: set[str] = set()
        for result in results:
            normalized_url = result["url"].rstrip("/").lower()
            if normalized_url in seen_urls:
                continue
            seen_urls.add(normalized_url)
            unique_results.append(result)
        results = unique_results[:6]
        sources = [{"title": result["title"], "url": result["url"]} for result in results]
        if not results:
            return system_prompt, sources
        context = "\n\n".join(
            f"Title: {result['title']}\nURL: {result['url']}\nSnippet: {result['snippet']}" for result in results
        )
        return f"{system_prompt}\n\nLive web results for the user's query:\n{context}", sources
    except Exception:
        traceback.print_exc()
        return system_prompt, []


class UrlRequest(BaseModel):
    url: HttpUrl


class OwnerAuthRequest(BaseModel):
    passkey: str

class OwnerAuthResponse(BaseModel):
    authenticated: bool
    session_id: str
    expires_in_seconds: int


def _owner_auth_allowed(client_key: str) -> bool:
    now = time.time()
    attempts = [stamp for stamp in _owner_auth_failures.get(client_key, []) if now - stamp < OWNER_AUTH_WINDOW_SECONDS]
    _owner_auth_failures[client_key] = attempts
    return len(attempts) < OWNER_AUTH_MAX_ATTEMPTS


@app.post("/auth/owner", response_model=OwnerAuthResponse)
async def auth_owner(req: OwnerAuthRequest, request: Request):
    """Issue a short-lived owner token after a rate-limited server-side passkey check."""
    if not OWNER_PASSKEY:
        raise HTTPException(status_code=503, detail="Owner access is not configured.")
    client_key = request.client.host if request.client else "unknown"
    if not _owner_auth_allowed(client_key):
        raise HTTPException(status_code=429, detail="Too many owner sign-in attempts. Try again later.")
    if not hmac.compare_digest(req.passkey, OWNER_PASSKEY):
        _owner_auth_failures.setdefault(client_key, []).append(time.time())
        raise HTTPException(status_code=403, detail="Invalid passkey")
    _owner_auth_failures.pop(client_key, None)
    return OwnerAuthResponse(
        authenticated=True,
        session_id=create_owner_session(),
        expires_in_seconds=OWNER_SESSION_TTL_SECONDS,
    )


@app.post("/fetch-url")
async def fetch_url(request: UrlRequest):
    if not is_safe_url(str(request.url)):
        raise HTTPException(status_code=400, detail="That URL can't be fetched.")

    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36"
        }
        resp = requests.get(str(request.url), headers=headers, timeout=20, stream=True)
        resp.raise_for_status()

        # Read up to 3MB of decoded content without touching internal resp._content
        raw_bytes = resp.raw.read(3 * 1024 * 1024 + 1, decode_content=True)
        if len(raw_bytes) > 3 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="Page is too large to read.")
        page_html = raw_bytes.decode(resp.encoding or "utf-8", errors="ignore")

        soup = BeautifulSoup(page_html, "html.parser")

        for tag in soup(["script", "style", "noscript", "svg", "canvas"]):
            tag.decompose()

        title = ""
        if soup.title and soup.title.string:
            title = soup.title.string.strip()

        text = " ".join(soup.stripped_strings)
        text = " ".join(text.split())

        if len(text) > 20000:
            text = text[:20000]

        if not text:
            raise HTTPException(status_code=400, detail="No readable text found on the page.")

        return {
            "url": str(request.url),
            "title": title,
            "content": text
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Could not fetch URL: {str(e)}")


@app.post("/upload-file")
async def upload_file(file: UploadFile = File(...)):
    file_type = get_file_type(file.filename)

    if file_type == 'unsupported':
        raise HTTPException(status_code=400, detail="Unsupported file type. Supported: PDF, DOCX, TXT, MD, PNG, JPG, JPEG, WEBP")

    ext = file.filename.lower().split('.')[-1] if '.' in file.filename else ''
    safe_name = f"{uuid.uuid4().hex}.{ext}" if ext else uuid.uuid4().hex
    file_path = os.path.join(UPLOAD_DIR, safe_name)

    try:
        size = 0
        with open(file_path, "wb") as buffer:
            while chunk := file.file.read(1024 * 1024):
                size += len(chunk)
                if size > MAX_UPLOAD_BYTES:
                    raise HTTPException(status_code=400, detail="File is too large. Max 10MB.")
                buffer.write(chunk)

        result = process_file(file_path, file.filename)

        if result['type'] == 'image':
            return {
                "filename": file.filename,
                "type": "image",
                "base64": result['content'],
                "mime": result['metadata'].get('mime_type', 'image/png')
            }
        elif result['type'] == 'text':
            return {
                "filename": file.filename,
                "type": "text",
                "content": result['content']
            }
        else:
            return {
                "filename": file.filename,
                "type": "error",
                "content": result.get('content', 'Could not read file')
            }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to process file: {str(e)}")
    finally:
        if os.path.exists(file_path):
            os.remove(file_path)


@app.get("/usage")
async def usage(x_guest_id: str = Header(default=""), x_mavis_session: str = Header(default="")):
    if validate_owner_session(x_mavis_session):
        return {"mode": "owner", "used": 0, "limit": None, "remaining": None}
    guest_id = _validate_guest_id(x_guest_id)
    return {"mode": "demo", **_guest_usage(guest_id)}


@app.post("/chat")
async def chat(request: ChatRequest, x_guest_id: str = Header(default="")):
    _validate_chat_request(request)
    guest_id = _validate_guest_id(x_guest_id)
    conv_id, is_owner = _conversation_identity(request, guest_id)
    _enforce_public_rate_limit(guest_id, is_owner)
    if not is_owner and check_guest_limit(guest_id):
        return {
            "response": "You have reached the 10-message Mavis demo limit. Thanks for trying her.",
            "title": None,
            "conv_id": conv_id,
            "limit_reached": True,
            "usage": _guest_usage(guest_id),
        }

    if not request.incognito:
        memory.create_conversation(conv_id, "New conversation", "owner" if is_owner else "guest", guest_id)
    history = _conversation_history(request, conv_id)
    system_prompt = _system_prompt_for(request, is_owner)
    search_prompt, sources = _search_context_for(request, system_prompt)
    try:
        assistant_message, provider = _call_text_with_fallback(
            [{"role": "system", "content": search_prompt}] + history + [{"role": "user", "content": request.message}],
            search_prompt,
        )
        assistant_message = _clean_assistant_response(assistant_message)
    except Exception:
        traceback.print_exc()
        raise HTTPException(status_code=503, detail="Mavis is temporarily unavailable. Please try again shortly.")

    title = None
    if not request.incognito:
        first_turn = not history
        memory.add_message(conv_id, "user", request.message)
        memory.add_message(conv_id, "assistant", assistant_message)
        if first_turn:
            title = _generate_title(request.message, assistant_message)
            memory.update_conversation_title(conv_id, title)

    usage_data = None if is_owner else _guest_usage(guest_id)
    if not is_owner:
        increment_guest_count(guest_id)
        usage_data = _guest_usage(guest_id)
    return {
        "response": assistant_message,
        "title": title,
        "conv_id": conv_id,
        "sources": sources,
        "provider": provider,
        "usage": usage_data,
    }


@app.post("/chat/stream")
async def chat_stream(request: ChatRequest, x_guest_id: str = Header(default="")):
    from fastapi.responses import StreamingResponse

    _validate_chat_request(request)
    guest_id = _validate_guest_id(x_guest_id)
    conv_id, is_owner = _conversation_identity(request, guest_id)
    _enforce_public_rate_limit(guest_id, is_owner)

    async def event_stream():
        if not is_owner and check_guest_limit(guest_id):
            yield "data: " + json.dumps({"type": "text", "content": "You have reached the 10-message Mavis demo limit. Thanks for trying her."}) + "\n\n"
            yield "data: " + json.dumps({"type": "done", "limit_reached": True, "usage": _guest_usage(guest_id)}) + "\n\n"
            return

        if not request.incognito:
            memory.create_conversation(conv_id, "New conversation", "owner" if is_owner else "guest", guest_id)
        history = _conversation_history(request, conv_id)
        system_prompt = _system_prompt_for(request, is_owner)
        search_prompt, sources = _search_context_for(request, system_prompt)
        try:
            response, provider = _call_text_with_fallback(
                [{"role": "system", "content": search_prompt}] + history + [{"role": "user", "content": request.message}],
                search_prompt,
            )
            response = _clean_assistant_response(response)
        except Exception:
            traceback.print_exc()
            yield "data: " + json.dumps({"type": "error", "content": "Mavis is temporarily unavailable. Please try again shortly."}) + "\n\n"
            return

        if response:
            yield "data: " + json.dumps({"type": "text", "content": response}) + "\n\n"
        if sources:
            yield "data: " + json.dumps({"type": "sources", "content": sources}) + "\n\n"

        title = None
        if not request.incognito:
            first_turn = not history
            memory.add_message(conv_id, "user", request.message)
            memory.add_message(conv_id, "assistant", response)
            if first_turn:
                title = _generate_title(request.message, response)
                memory.update_conversation_title(conv_id, title)

        usage_data = None
        if not is_owner:
            increment_guest_count(guest_id)
            usage_data = _guest_usage(guest_id)
        yield "data: " + json.dumps({"type": "done", "title": title, "conv_id": conv_id, "provider": provider, "usage": usage_data}) + "\n\n"

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


class CreateFileRequest(BaseModel):
    content: str
    filename: str = "document"
    file_type: str = "docx"  # docx, txt, md, html

@app.post("/create-file")
async def create_file(req: CreateFileRequest):
    """Generate a downloadable file from content."""
    content = req.content
    filename = req.filename
    file_type = req.file_type

    if file_type == "docx":
        doc = DocxDocument()
        # Parse markdown-style formatting
        for line in content.split('\n'):
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('```'):
                pass  # skip code fence markers
            else:
                if line.strip():
                    doc.add_paragraph(line)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        base64_data = base64.b64encode(buf.read()).decode()
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = "docx"
    elif file_type == "txt":
        base64_data = base64.b64encode(content.encode('utf-8')).decode()
        mime = "text/plain"
        ext = "txt"
    elif file_type == "md":
        base64_data = base64.b64encode(content.encode('utf-8')).decode()
        mime = "text/markdown"
        ext = "md"
    elif file_type == "html":
        base64_data = base64.b64encode(content.encode('utf-8')).decode()
        mime = "text/html"
        ext = "html"
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_type}")

    return {
      "success": True,
      "filename": f"{filename}.{ext}",
      "base64": base64_data,
      "mime": mime,
      "file_ready": True,
      "message": f"Your {ext.toUpperCase()} file is ready"
    }


def _attachment_request(
    message: str,
    session_id: str,
    owner_session: str,
    incognito: bool,
    persona: str,
    history_json: str,
) -> ChatRequest:
    try:
        parsed_history = json.loads(history_json or "[]")
    except (TypeError, ValueError):
        parsed_history = []
    safe_history = [
        item
        for item in parsed_history
        if isinstance(item, dict)
        and isinstance(item.get("role"), str)
        and isinstance(item.get("content"), str)
    ]
    request = ChatRequest(
        message=message or "Please analyze this attachment.",
        session_id=session_id,
        owner_session=owner_session,
        incognito=incognito,
        persona=persona,
        web_search=False,
        history=safe_history,
    )
    _validate_chat_request(request)
    return request


@app.post("/chat-with-file")
async def chat_with_file(
    message: str = Form(""),
    session_id: str = Form(...),
    owner_session: str = Form(""),
    incognito: bool = Form(False),
    persona: str = Form("default"),
    history: str = Form("[]"),
    filename: str = Form(...),
    file_content: str = Form(""),
    x_guest_id: str = Header(default=""),
):
    request = _attachment_request(message, session_id, owner_session, incognito, persona, history)
    guest_id = _validate_guest_id(x_guest_id)
    if not filename or len(filename) > 200 or len(file_content) > 50_000:
        raise HTTPException(status_code=400, detail="The attachment is invalid or too large to analyze.")
    conv_id, is_owner = _conversation_identity(request, guest_id)
    _enforce_public_rate_limit(guest_id, is_owner)
    if not is_owner and check_guest_limit(guest_id):
        return {"response": "You have reached the 10-message Mavis demo limit. Thanks for trying her.", "limit_reached": True, "usage": _guest_usage(guest_id)}

    if not request.incognito:
        memory.create_conversation(conv_id, "New conversation", "owner" if is_owner else "guest", guest_id)
    history = _conversation_history(request, conv_id)
    system_prompt = _system_prompt_for(request, is_owner)
    prompt = f"{request.message}\n\nAttached file ({filename}):\n{file_content}"
    try:
        response, provider = _call_text_with_fallback(
            [{"role": "system", "content": system_prompt}] + history + [{"role": "user", "content": prompt}],
            system_prompt,
        )
    except Exception:
        traceback.print_exc()
        raise HTTPException(status_code=503, detail="Mavis could not analyze that attachment right now.")

    title = None
    if not request.incognito:
        first_turn = not history
        memory.add_message(conv_id, "user", f"{request.message} [attached {filename}]")
        memory.add_message(conv_id, "assistant", response)
        if first_turn:
            title = _generate_title(request.message, response)
            memory.update_conversation_title(conv_id, title)
    usage_data = None
    if not is_owner:
        increment_guest_count(guest_id)
        usage_data = _guest_usage(guest_id)
    return {"response": response, "title": title, "conv_id": conv_id, "provider": provider, "usage": usage_data}


def _call_image_with_fallback(prompt: str, system_prompt: str, base64_image: str, mime: str) -> tuple[str, str]:
    if len(base64_image) > 8_000_000:
        raise HTTPException(status_code=400, detail="Images are limited to 6 MB.")
    if GEMINI_API_KEY:
        try:
            response = requests.post(
                f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent?key={GEMINI_API_KEY}",
                json={
                    "system_instruction": {"parts": [{"text": system_prompt}]},
                    "contents": [{"role": "user", "parts": [{"text": prompt}, {"inline_data": {"mime_type": mime, "data": base64_image}}]}],
                },
                timeout=60,
            )
            response.raise_for_status()
            data = response.json()
            parts = (data.get("candidates") or [{}])[0].get("content", {}).get("parts", [])
            answer = "".join(part.get("text", "") for part in parts if isinstance(part, dict))
            if answer:
                return answer, "gemini"
            raise RuntimeError("Gemini returned no usable image response")
        except Exception as error:
            if not _is_retryable_provider_error(error):
                raise HTTPException(status_code=503, detail="Mavis could not analyze that image right now.") from error
    try:
        response = get_groq_client().chat.completions.create(
            model=os.environ.get("GROQ_VISION_MODEL", "meta-llama/llama-4-scout-17b-16e-instruct"),
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": [{"type": "text", "text": prompt}, {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{base64_image}"}}]},
            ],
            max_tokens=1024,
        )
        return response.choices[0].message.content or "", "groq"
    except Exception as error:
        raise HTTPException(status_code=503, detail="Mavis could not analyze that image right now.") from error


@app.post("/chat-with-image")
async def chat_with_image(
    message: str = Form(""),
    session_id: str = Form(...),
    owner_session: str = Form(""),
    incognito: bool = Form(False),
    persona: str = Form("default"),
    history: str = Form("[]"),
    base64_image: str = Form(...),
    mime: str = Form(...),
    x_guest_id: str = Header(default=""),
):
    request = _attachment_request(message, session_id, owner_session, incognito, persona, history)
    guest_id = _validate_guest_id(x_guest_id)
    if mime not in {"image/jpeg", "image/png", "image/webp", "image/gif"}:
        raise HTTPException(status_code=400, detail="Unsupported image type.")
    conv_id, is_owner = _conversation_identity(request, guest_id)
    _enforce_public_rate_limit(guest_id, is_owner)
    if not is_owner and check_guest_limit(guest_id):
        return {"response": "You have reached the 10-message Mavis demo limit. Thanks for trying her.", "limit_reached": True, "usage": _guest_usage(guest_id)}

    if not request.incognito:
        memory.create_conversation(conv_id, "New conversation", "owner" if is_owner else "guest", guest_id)
    history = _conversation_history(request, conv_id)
    system_prompt = _system_prompt_for(request, is_owner)
    response, provider = _call_image_with_fallback(request.message, system_prompt, base64_image, mime)
    response = _clean_assistant_response(response)
    title = None
    if not request.incognito:
        first_turn = not history
        memory.add_message(conv_id, "user", f"{request.message} [shared an image]")
        memory.add_message(conv_id, "assistant", response)
        if first_turn:
            title = _generate_title(request.message, response)
            memory.update_conversation_title(conv_id, title)
    usage_data = None
    if not is_owner:
        increment_guest_count(guest_id)
        usage_data = _guest_usage(guest_id)
    return {"response": response, "title": title, "conv_id": conv_id, "provider": provider, "usage": usage_data}


@app.get("/conversations")
async def list_conversations(x_guest_id: str = Header(default=""), x_mavis_session: str = Header(default="")):
    if validate_owner_session(x_mavis_session):
        return {"conversations": memory.get_all_owner_conversations()}
    return {"conversations": memory.get_all_conversations_for_guest(_validate_guest_id(x_guest_id))}


@app.get("/conversations/{conv_id}/messages")
async def get_messages(conv_id: str, x_guest_id: str = Header(default=""), x_mavis_session: str = Header(default="")):
    owner_access = validate_owner_session(x_mavis_session)
    guest_access = memory.conversation_belongs_to_guest(conv_id, _validate_guest_id(x_guest_id))
    if not ((owner_access and memory.conversation_is_owner(conv_id)) or guest_access):
        raise HTTPException(status_code=404, detail="Conversation not found")
    return {"messages": memory.get_conversation_messages(conv_id, MAX_HISTORY_MESSAGES)}


@app.delete("/conversations/{conv_id}")
async def delete_conv(conv_id: str, x_guest_id: str = Header(default=""), x_mavis_session: str = Header(default="")):
    owner_access = validate_owner_session(x_mavis_session)
    guest_access = memory.conversation_belongs_to_guest(conv_id, _validate_guest_id(x_guest_id))
    if not ((owner_access and memory.conversation_is_owner(conv_id)) or guest_access):
        raise HTTPException(status_code=404, detail="Conversation not found")
    memory.delete_conversation(conv_id)
    return {"status": "deleted"}


CODING_SYSTEM_PROMPT = """
You are Mavis in Coding Mode: a careful senior software engineer working inside one
owner-approved project workspace. You do not have shell access and must never claim to
have applied, tested, or deployed a change. Analyze only the source files provided.

Return valid JSON only—no Markdown fence or prose before/after it—with this exact shape:
{
  "summary": "one concise sentence",
  "answer": "a concise helpful answer when the user asks to understand or review code; otherwise an empty string",
  "plan": ["short ordered step"],
  "questions": ["only questions that block a safe change"],
  "changes": [
    {
      "path": "selected/project/file.ext",
      "operation": "replace",
      "find": "an exact unique excerpt copied from the supplied file",
      "replace": "the replacement text",
      "explanation": "why this focused edit is needed"
    }
  ],
  "verification": ["frontend_build"]
}

Rules:
- Make no more than six focused changes. Use only the supplied selected paths.
- Initial Coding Mode supports only operation "replace" in existing selected files.
  A replace change must contain exact text that occurs once in the source file.
- If the user asks only to understand, review, or explain the selected files, return no changes and provide the explanation in answer.
- If requirements are unclear or the requested change is unsafe, return no changes and
  use questions to ask for clarification.
- Never include secrets, environment variable values, dependency lockfiles, node_modules,
  or generated build files.
- Available verification IDs: frontend_build, backend_tests, deployment_tests.
- Prefer minimal, maintainable edits that preserve the existing design and behavior.
"""

TEMPORARY_PROJECT_SYSTEM_PROMPT = """
You are Mavis in Temporary Project Mode: a careful senior software engineer reviewing a
small, owner-uploaded project stored in an isolated temporary workspace. You do not have
shell access and must never claim to have applied, tested, or deployed a change. Analyze
only the supplied uploaded source files.

Return valid JSON only—no Markdown fence or prose before/after it—with this exact shape:
{
  "summary": "one concise sentence",
  "answer": "a concise helpful answer when the user asks to understand or review code; otherwise an empty string",
  "plan": ["short ordered step"],
  "questions": ["only questions that block a safe change"],
  "changes": [
    {
      "path": "selected/uploaded/file.ext",
      "operation": "replace",
      "find": "an exact unique excerpt copied from the supplied file",
      "replace": "the replacement text",
      "explanation": "why this focused edit is needed"
    }
  ],
  "verification": ["project_scan"]
}

Rules:
- Make no more than six focused changes. Use only the supplied selected paths.
- Only operation "replace" is supported; never create, delete, rename, or move files.
- A replace change must contain exact text that occurs once in the supplied file.
- If requirements are unclear or a change is unsafe, return no changes and use questions.
- Never request, reveal, or include secrets or environment variable values.
- Available verification IDs: project_scan, json_validate. These checks never execute uploaded code.
- Prefer minimal, maintainable edits that preserve the existing design and behavior.
"""


class CodingTaskRequest(BaseModel):
    message: str = Field(min_length=1, max_length=MAX_MESSAGE_CHARACTERS)
    session_id: str = Field(min_length=8, max_length=128)
    owner_session: str = Field(min_length=1, max_length=256)
    files: list[str] = Field(min_length=1, max_length=coding_workspace.MAX_SELECTED_FILES)
    history: list[dict[str, str]] = Field(default_factory=list)


class CodingProposalAction(BaseModel):
    proposal_id: str = Field(min_length=8, max_length=64)
    confirm: bool = False


class TemporaryProjectTaskRequest(CodingTaskRequest):
    project_id: str = Field(min_length=37, max_length=37)


class TemporaryProjectAction(CodingProposalAction):
    project_id: str = Field(min_length=37, max_length=37)


class CodingVerificationRequest(BaseModel):
    proposal_id: str = Field(min_length=8, max_length=64)
    command: str = Field(min_length=1, max_length=64)


class TemporaryProjectVerificationRequest(CodingVerificationRequest):
    project_id: str = Field(min_length=37, max_length=37)


def _coding_owner_key(owner_session: str) -> str:
    return hashlib.sha256(owner_session.encode("utf-8")).hexdigest()


def _require_coding_owner(owner_session: str) -> str:
    if not validate_owner_session(owner_session):
        raise HTTPException(status_code=401, detail="Coding Mode requires owner access.")
    if not coding_workspace.workspace_enabled():
        raise HTTPException(
            status_code=503,
            detail="Coding Mode is not enabled on this deployment yet. Set CODE_WORKSPACE_ENABLED=true only on the private owner backend.",
        )
    return _coding_owner_key(owner_session)


def _require_temporary_project_owner(owner_session: str) -> str:
    if not validate_owner_session(owner_session):
        raise HTTPException(status_code=401, detail="Temporary Project Mode requires owner access.")
    if not temporary_project_workspace.workspace_enabled():
        raise HTTPException(
            status_code=503,
            detail="Temporary Project Mode is not enabled on this deployment yet.",
        )
    return _coding_owner_key(owner_session)


def _prune_coding_proposals() -> None:
    cutoff = time.time() - CODE_PROPOSAL_TTL_SECONDS
    expired = [proposal_id for proposal_id, proposal in _coding_proposals.items() if proposal["created_at"] < cutoff]
    for proposal_id in expired:
        proposal = _coding_proposals.pop(proposal_id)
        checkpoint_id = proposal.get("checkpoint_id")
        if checkpoint_id and proposal.get("workspace") == "mavis":
            coding_workspace.cleanup_checkpoint(str(checkpoint_id))


def _prune_temporary_projects() -> None:
    expired_ids = set(temporary_project_workspace.cleanup_expired_projects(TEMP_PROJECT_TTL_SECONDS))
    cutoff = time.time() - TEMP_PROJECT_TTL_SECONDS
    for project_id, project in list(_temporary_projects.items()):
        if project_id in expired_ids or project["created_at"] < cutoff:
            _temporary_projects.pop(project_id, None)


def _temporary_project_or_404(project_id: str, owner_key: str) -> dict[str, Any]:
    _prune_temporary_projects()
    project = _temporary_projects.get(project_id)
    if not project or project["owner_key"] != owner_key:
        raise HTTPException(status_code=404, detail="That temporary project has expired or is unavailable.")
    try:
        temporary_project_workspace.list_project_files(project_id)
    except temporary_project_workspace.TemporaryProjectError as error:
        _temporary_projects.pop(project_id, None)
        raise HTTPException(status_code=404, detail="That temporary project has expired or is unavailable.") from error
    return project


def _proposal_or_404(proposal_id: str, owner_key: str) -> dict[str, Any]:
    _prune_coding_proposals()
    proposal = _coding_proposals.get(proposal_id)
    if not proposal or proposal["owner_key"] != owner_key:
        raise HTTPException(status_code=404, detail="That coding proposal is unavailable or has expired.")
    return proposal


def _coding_failure_detail(error: Exception) -> str:
    details: list[str] = []
    current: BaseException | None = error
    while current is not None and len(details) < 4:
        details.append(str(current).lower())
        current = current.__cause__ or current.__context__
    detail = " ".join(details)
    if any(token in detail for token in ("413", "request_too_large", "request entity too large")):
        return "The selected files contain too much code for one plan. Choose one to three smaller, focused files."
    if any(token in detail for token in ("429", "rate limit", "quota", "resource_exhausted")):
        return "The coding provider is temporarily rate-limited. Wait a minute, then retry with fewer selected files."
    if any(token in detail for token in ("timeout", "timed out", "504")):
        return "The coding provider took too long to prepare a plan. Retry with fewer selected files."
    return "Mavis could not prepare a coding plan right now. Retry once; if it persists, choose fewer files and try again."


def _clean_json_response(response: str) -> dict[str, Any]:
    raw = response.strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?\\s*|\\s*```$", "", raw, flags=re.IGNORECASE)
    try:
        parsed = json.loads(raw)
    except (TypeError, ValueError):
        start, end = raw.find("{"), raw.rfind("}")
        if start < 0 or end <= start:
            raise ValueError("Mavis returned an invalid coding proposal. Please try again.")
        try:
            parsed = json.loads(raw[start : end + 1])
        except (TypeError, ValueError) as error:
            raise ValueError("Mavis returned an invalid coding proposal. Please try again.") from error
    if not isinstance(parsed, dict):
        raise ValueError("Mavis returned an invalid coding proposal. Please try again.")
    return parsed


def _normalize_coding_proposal(
    raw: dict[str, Any],
    context: list[dict[str, str]],
    allowed_verification: set[str] | None = None,
) -> dict[str, Any]:
    selected = {item["path"]: item["content"] for item in context}
    summary = str(raw.get("summary") or "Mavis reviewed the selected project files.").strip()[:600]
    answer = str(raw.get("answer") or "").strip()[:6_000]
    plan = [str(item).strip()[:300] for item in raw.get("plan", []) if str(item).strip()][:8]
    questions = [str(item).strip()[:400] for item in raw.get("questions", []) if str(item).strip()][:5]
    verification = [str(item).strip() for item in raw.get("verification", []) if str(item).strip()]
    allowed_checks = allowed_verification or {"frontend_build", "backend_tests", "deployment_tests"}
    verification = [item for item in verification if item in allowed_checks][:3]

    changes: list[dict[str, str]] = []
    diffs: list[dict[str, str]] = []
    for candidate in raw.get("changes", [])[:6]:
        if not isinstance(candidate, dict):
            continue
        path = str(candidate.get("path") or "").strip()
        operation = str(candidate.get("operation") or "replace").strip()
        find = str(candidate.get("find") or "")
        replace = str(candidate.get("replace") or "")
        explanation = str(candidate.get("explanation") or "Focused code update.").strip()[:500]
        if path not in selected or operation != "replace" or not find:
            continue
        if len(find.encode("utf-8")) > coding_workspace.MAX_FILE_BYTES or len(replace.encode("utf-8")) > coding_workspace.MAX_FILE_BYTES:
            continue
        before = selected[path]
        if before.count(find) != 1:
            continue
        after = before.replace(find, replace, 1)
        diff = "".join(
            difflib.unified_diff(
                before.splitlines(keepends=True),
                after.splitlines(keepends=True),
                fromfile=f"a/{path}",
                tofile=f"b/{path}",
            )
        )[:12_000]
        changes.append({"path": path, "operation": operation, "find": find, "replace": replace, "explanation": explanation})
        diffs.append({"path": path, "diff": diff})

    if not plan:
        plan = ["Review the selected project context", "Propose only the smallest safe edits", "Run the recommended verification after approval"]
    if not changes and not questions and not answer:
        questions = ["I could not produce a safe exact-match patch from the selected files. Please refresh the workspace and try again."]
    return {
        "summary": summary,
        "answer": answer,
        "plan": plan,
        "questions": questions,
        "changes": changes,
        "verification": verification,
        "diffs": diffs,
    }


def _coding_history(history: list[dict[str, str]]) -> list[dict[str, str]]:
    normalized: list[dict[str, str]] = []
    for item in history[-8:]:
        if not isinstance(item, dict):
            continue
        role = item.get("role")
        content = item.get("content")
        if role in {"user", "assistant"} and isinstance(content, str) and content.strip():
            normalized.append({"role": role, "content": content.strip()[:4_000]})
    return normalized


def _public_coding_proposal(record: dict[str, Any]) -> dict[str, Any]:
    public_record = {
        key: value
        for key, value in record.items()
        if key not in {"owner_key", "created_at", "changes"}
    }
    public_record["proposed_changes"] = [
        {
            "path": change["path"],
            "operation": change["operation"],
            "explanation": change["explanation"],
        }
        for change in record["changes"]
    ]
    return public_record


@app.post("/temporary-projects")
async def create_temporary_project(
    files: list[UploadFile] = File(...),
    x_mavis_session: str = Header(default=""),
):
    owner_key = _require_temporary_project_owner(x_mavis_session)
    _prune_temporary_projects()
    uploads: list[tuple[str, bytes]] = []
    for upload in files:
        filename = (upload.filename or "").strip()
        if not filename:
            raise HTTPException(status_code=400, detail="Every uploaded file needs a filename.")
        raw = await upload.read(temporary_project_workspace.MAX_FILE_BYTES + 1)
        uploads.append((filename, raw))
    try:
        project = temporary_project_workspace.create_project(uploads)
    except temporary_project_workspace.TemporaryProjectError as error:
        raise HTTPException(status_code=400, detail=str(error)) from error
    project_id = str(project["project_id"])
    _temporary_projects[project_id] = {"owner_key": owner_key, "created_at": time.time()}
    return {**project, "expires_in_seconds": TEMP_PROJECT_TTL_SECONDS}


@app.get("/temporary-projects/{project_id}")
async def temporary_project_status(project_id: str, x_mavis_session: str = Header(default="")):
    owner_key = _require_temporary_project_owner(x_mavis_session)
    _temporary_project_or_404(project_id, owner_key)
    try:
        files = temporary_project_workspace.list_project_files(project_id)
    except temporary_project_workspace.TemporaryProjectError as error:
        raise HTTPException(status_code=404, detail=str(error)) from error
    return {"project_id": project_id, "files": files, "expires_in_seconds": TEMP_PROJECT_TTL_SECONDS}


@app.get("/temporary-projects/{project_id}/file")
async def temporary_project_file(project_id: str, path: str, x_mavis_session: str = Header(default="")):
    owner_key = _require_temporary_project_owner(x_mavis_session)
    _temporary_project_or_404(project_id, owner_key)
    try:
        return temporary_project_workspace.read_project_file(project_id, path)
    except temporary_project_workspace.TemporaryProjectError as error:
        raise HTTPException(status_code=400, detail=str(error)) from error


@app.post("/temporary-projects/{project_id}/propose")
async def propose_temporary_project_change(project_id: str, request: TemporaryProjectTaskRequest):
    owner_key = _require_temporary_project_owner(request.owner_session)
    _temporary_project_or_404(project_id, owner_key)
    if project_id != request.project_id:
        raise HTTPException(status_code=400, detail="The temporary project identifier does not match the request.")
    if not re.fullmatch(r"[A-Za-z0-9-]{8,128}", request.session_id):
        raise HTTPException(status_code=400, detail="Invalid conversation identifier.")
    task = request.message.strip()
    if not task:
        raise HTTPException(status_code=400, detail="Describe the coding task first.")
    try:
        context = temporary_project_workspace.read_project_context(project_id, request.files)
    except temporary_project_workspace.TemporaryProjectError as error:
        raise HTTPException(status_code=400, detail=str(error)) from error

    file_context = "\n\n".join(
        f"--- UPLOADED FILE: {item['path']} ---\n{item['content']}\n--- END FILE ---" for item in context
    )
    prompt = f"User's temporary-project coding task:\n{task}\n\nSelected uploaded files:\n{file_context}"
    try:
        model_response, provider = _call_text_with_fallback(
            [{"role": "system", "content": TEMPORARY_PROJECT_SYSTEM_PROMPT}]
            + _coding_history(request.history)
            + [{"role": "user", "content": prompt}],
            TEMPORARY_PROJECT_SYSTEM_PROMPT,
            temperature=0.2,
            max_tokens=3_500,
            json_mode=True,
        )
        proposal = _normalize_coding_proposal(
            _clean_json_response(model_response),
            context,
            {"project_scan", "json_validate"},
        )
    except ValueError as error:
        raise HTTPException(status_code=502, detail=str(error)) from error
    except Exception as error:
        traceback.print_exc()
        raise HTTPException(status_code=503, detail=_coding_failure_detail(error)) from error

    proposal_id = uuid.uuid4().hex
    record = {
        **proposal,
        "proposal_id": proposal_id,
        "owner_key": owner_key,
        "created_at": time.time(),
        "status": "pending",
        "checkpoint_id": None,
        "changed_files": [],
        "provider": provider,
        "workspace": "temporary",
        "project_id": project_id,
    }
    _coding_proposals[proposal_id] = record
    return _public_coding_proposal(record)


@app.post("/temporary-projects/{project_id}/apply")
async def apply_temporary_project_change(
    project_id: str,
    request: TemporaryProjectAction,
    x_mavis_session: str = Header(default=""),
):
    owner_key = _require_temporary_project_owner(x_mavis_session)
    _temporary_project_or_404(project_id, owner_key)
    if project_id != request.project_id:
        raise HTTPException(status_code=400, detail="The temporary project identifier does not match the request.")
    if not request.confirm:
        raise HTTPException(status_code=400, detail="Confirm the proposal before Mavis can apply it.")
    proposal = _proposal_or_404(request.proposal_id, owner_key)
    if proposal.get("workspace") != "temporary" or proposal.get("project_id") != project_id:
        raise HTTPException(status_code=404, detail="That proposal does not belong to this temporary project.")
    if proposal["status"] != "pending":
        raise HTTPException(status_code=409, detail="This coding proposal was already handled. Create a new proposal to continue.")
    try:
        result = temporary_project_workspace.apply_changes(project_id, proposal["changes"])
    except temporary_project_workspace.TemporaryProjectError as error:
        raise HTTPException(status_code=409, detail=str(error)) from error
    proposal["status"] = "applied"
    proposal["checkpoint_id"] = result["checkpoint_id"]
    proposal["changed_files"] = result["changed_files"]
    return {"status": "applied", **result}


@app.post("/temporary-projects/{project_id}/verify")
async def verify_temporary_project_change(
    project_id: str,
    request: TemporaryProjectVerificationRequest,
    x_mavis_session: str = Header(default=""),
):
    owner_key = _require_temporary_project_owner(x_mavis_session)
    _temporary_project_or_404(project_id, owner_key)
    if project_id != request.project_id:
        raise HTTPException(status_code=400, detail="The temporary project identifier does not match the request.")
    proposal = _proposal_or_404(request.proposal_id, owner_key)
    if proposal.get("workspace") != "temporary" or proposal.get("project_id") != project_id:
        raise HTTPException(status_code=404, detail="That proposal does not belong to this temporary project.")
    if proposal["status"] != "applied":
        raise HTTPException(status_code=409, detail="Apply the proposal before running its verification checks.")
    if request.command not in proposal["verification"]:
        raise HTTPException(status_code=400, detail="That check was not recommended for this proposal.")
    try:
        result = temporary_project_workspace.verify_project(project_id, request.command)
    except temporary_project_workspace.TemporaryProjectError as error:
        raise HTTPException(status_code=400, detail=str(error)) from error
    proposal.setdefault("verification_results", []).append(result)
    return result


@app.post("/temporary-projects/{project_id}/rollback")
async def rollback_temporary_project_change(
    project_id: str,
    request: TemporaryProjectAction,
    x_mavis_session: str = Header(default=""),
):
    owner_key = _require_temporary_project_owner(x_mavis_session)
    _temporary_project_or_404(project_id, owner_key)
    if project_id != request.project_id:
        raise HTTPException(status_code=400, detail="The temporary project identifier does not match the request.")
    if not request.confirm:
        raise HTTPException(status_code=400, detail="Confirm the rollback before restoring the uploaded files.")
    proposal = _proposal_or_404(request.proposal_id, owner_key)
    if proposal.get("workspace") != "temporary" or proposal.get("project_id") != project_id:
        raise HTTPException(status_code=404, detail="That proposal does not belong to this temporary project.")
    if proposal["status"] != "applied" or not proposal.get("checkpoint_id"):
        raise HTTPException(status_code=409, detail="There is no applied coding proposal to undo.")
    try:
        restored = temporary_project_workspace.rollback_checkpoint(project_id, str(proposal["checkpoint_id"]), proposal["changed_files"])
    except temporary_project_workspace.TemporaryProjectError as error:
        raise HTTPException(status_code=409, detail=str(error)) from error
    proposal["status"] = "rolled_back"
    return {"status": "rolled_back", "restored_files": restored}


@app.get("/temporary-projects/{project_id}/download")
async def download_temporary_project(project_id: str, x_mavis_session: str = Header(default="")):
    owner_key = _require_temporary_project_owner(x_mavis_session)
    _temporary_project_or_404(project_id, owner_key)
    try:
        archive = temporary_project_workspace.export_project(project_id)
    except temporary_project_workspace.TemporaryProjectError as error:
        raise HTTPException(status_code=404, detail=str(error)) from error
    return Response(
        content=archive,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="mavis-edited-{project_id}.zip"'},
    )


@app.delete("/temporary-projects/{project_id}")
async def delete_temporary_project(project_id: str, confirm: bool = False, x_mavis_session: str = Header(default="")):
    owner_key = _require_temporary_project_owner(x_mavis_session)
    _temporary_project_or_404(project_id, owner_key)
    if not confirm:
        raise HTTPException(status_code=400, detail="Confirm before deleting the temporary project.")
    try:
        temporary_project_workspace.delete_project(project_id)
    except temporary_project_workspace.TemporaryProjectError as error:
        raise HTTPException(status_code=404, detail=str(error)) from error
    _temporary_projects.pop(project_id, None)
    for proposal_id, proposal in list(_coding_proposals.items()):
        if proposal.get("workspace") == "temporary" and proposal.get("project_id") == project_id:
            _coding_proposals.pop(proposal_id, None)
    return {"status": "deleted"}


@app.get("/coding/workspace")
async def coding_workspace_status(x_mavis_session: str = Header(default="")):
    _require_coding_owner(x_mavis_session)
    try:
        return {"enabled": True, "files": coding_workspace.list_workspace_files()}
    except coding_workspace.WorkspaceError as error:
        raise HTTPException(status_code=400, detail=str(error)) from error


@app.get("/coding/workspace/file")
async def coding_workspace_file(path: str, x_mavis_session: str = Header(default="")):
    _require_coding_owner(x_mavis_session)
    try:
        return coding_workspace.read_workspace_file(path)
    except coding_workspace.WorkspaceError as error:
        raise HTTPException(status_code=400, detail=str(error)) from error


@app.post("/coding/propose")
async def propose_coding_change(request: CodingTaskRequest):
    owner_key = _require_coding_owner(request.owner_session)
    if not re.fullmatch(r"[A-Za-z0-9-]{8,128}", request.session_id):
        raise HTTPException(status_code=400, detail="Invalid conversation identifier.")
    task = request.message.strip()
    if not task:
        raise HTTPException(status_code=400, detail="Describe the coding task first.")
    try:
        context = coding_workspace.read_workspace_context(request.files)
    except coding_workspace.WorkspaceError as error:
        raise HTTPException(status_code=400, detail=str(error)) from error

    file_context = "\n\n".join(
        f"--- FILE: {item['path']} ---\n{item['content']}\n--- END FILE ---" for item in context
    )
    prompt = f"User's coding task:\n{task}\n\nSelected workspace files:\n{file_context}"
    try:
        model_response, provider = _call_text_with_fallback(
            [{"role": "system", "content": CODING_SYSTEM_PROMPT}]
            + _coding_history(request.history)
            + [{"role": "user", "content": prompt}],
            CODING_SYSTEM_PROMPT,
            temperature=0.2,
            max_tokens=3_500,
            json_mode=True,
        )
        proposal = _normalize_coding_proposal(_clean_json_response(model_response), context)
    except ValueError as error:
        raise HTTPException(status_code=502, detail=str(error)) from error
    except Exception as error:
        traceback.print_exc()
        raise HTTPException(status_code=503, detail=_coding_failure_detail(error)) from error

    proposal_id = uuid.uuid4().hex
    record = {
        **proposal,
        "proposal_id": proposal_id,
        "owner_key": owner_key,
        "created_at": time.time(),
        "status": "pending",
        "checkpoint_id": None,
        "changed_files": [],
        "provider": provider,
    }
    _coding_proposals[proposal_id] = record
    public_record = {key: value for key, value in record.items() if key not in {"owner_key", "created_at", "changes"}}
    public_record["proposed_changes"] = [
        {"path": change["path"], "operation": change["operation"], "explanation": change["explanation"]}
        for change in record["changes"]
    ]
    return public_record


@app.post("/coding/apply")
async def apply_coding_change(request: CodingProposalAction, x_mavis_session: str = Header(default="")):
    owner_key = _require_coding_owner(x_mavis_session)
    if not request.confirm:
        raise HTTPException(status_code=400, detail="Confirm the proposal before Mavis can apply it.")
    proposal = _proposal_or_404(request.proposal_id, owner_key)
    if proposal["status"] != "pending":
        raise HTTPException(status_code=409, detail="This coding proposal was already handled. Create a new proposal to continue.")
    try:
        result = coding_workspace.apply_changes(proposal["changes"])
    except coding_workspace.WorkspaceError as error:
        raise HTTPException(status_code=409, detail=str(error)) from error
    proposal["status"] = "applied"
    proposal["checkpoint_id"] = result["checkpoint_id"]
    proposal["changed_files"] = result["changed_files"]
    return {"status": "applied", **result}


@app.post("/coding/verify")
async def verify_coding_change(request: CodingVerificationRequest, x_mavis_session: str = Header(default="")):
    owner_key = _require_coding_owner(x_mavis_session)
    proposal = _proposal_or_404(request.proposal_id, owner_key)
    if proposal["status"] != "applied":
        raise HTTPException(status_code=409, detail="Apply the proposal before running its verification checks.")
    if request.command not in proposal["verification"]:
        raise HTTPException(status_code=400, detail="That check was not recommended for this proposal.")
    try:
        result = coding_workspace.run_verification(request.command)
    except coding_workspace.WorkspaceError as error:
        raise HTTPException(status_code=400, detail=str(error)) from error
    proposal.setdefault("verification_results", []).append(result)
    return result


@app.post("/coding/rollback")
async def rollback_coding_change(request: CodingProposalAction, x_mavis_session: str = Header(default="")):
    owner_key = _require_coding_owner(x_mavis_session)
    if not request.confirm:
        raise HTTPException(status_code=400, detail="Confirm the rollback before restoring the previous files.")
    proposal = _proposal_or_404(request.proposal_id, owner_key)
    if proposal["status"] != "applied" or not proposal.get("checkpoint_id"):
        raise HTTPException(status_code=409, detail="There is no applied coding proposal to undo.")
    try:
        restored = coding_workspace.rollback_checkpoint(str(proposal["checkpoint_id"]), proposal["changed_files"])
    except coding_workspace.WorkspaceError as error:
        raise HTTPException(status_code=409, detail=str(error)) from error
    proposal["status"] = "rolled_back"
    return {"status": "rolled_back", "restored_files": restored}


@app.get("/health")
async def health():
    return {
        "status": "MAVIS is online",
        "model_provider": DEFAULT_MODEL_PROVIDER,
        "groq_configured": client is not None,
        "gemini_configured": bool(GEMINI_API_KEY),
        "guest_message_limit": GUEST_MESSAGE_LIMIT,
        "public_requests_per_minute": PUBLIC_REQUESTS_PER_MINUTE,
        "quota_storage": usage_store.backend_name(),
        "coding_mode_enabled": coding_workspace.workspace_enabled(),
        "temporary_project_mode_enabled": temporary_project_workspace.workspace_enabled(),
        "provider_metrics": dict(_provider_metrics),
    }
